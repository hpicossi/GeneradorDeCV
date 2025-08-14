import os
import re
import requests
from datetime import datetime
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
import json
import logging
import sqlite3
import argparse
import csv
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import requests
from bs4 import BeautifulSoup
import time
from urllib.parse import urljoin, urlparse
from typing import Optional, Dict, Any, List, Tuple

# Intentar cargar python-dotenv (opcional)
try:
    from dotenv import load_dotenv
    load_dotenv()  # Cargar variables de entorno desde .env
    DOTENV_AVAILABLE = True
except ImportError:
    DOTENV_AVAILABLE = False
    logging.warning("python-dotenv no está instalado. Para usar .env: pip install python-dotenv")

# Configurar logging con encoding UTF-8
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('cv_generator.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

class CVGeneratorError(Exception):
    """Excepción personalizada para errores del generador"""
    pass

class ConfigurationError(CVGeneratorError):
    """Error de configuración"""
    pass

class FileProcessingError(CVGeneratorError):
    """Error procesando archivos"""
    pass

class GeneradorCVInteligente:
    def __init__(self, config_path="config.json"):
        # Cargar configuración
        try:
            self.config = self.cargar_configuracion(config_path)
            self.cv_base_path = self.config['configuracion_general']['cv_base_path']
            self.carpeta_salida = self.config['configuracion_general']['carpeta_salida']
            self.umbral_fit = self.config['configuracion_general']['umbral_fit']
            self.perfil_tecnico = self.config['perfil_tecnico']
            logging.info(f"✅ Configuración cargada desde {config_path}")
        except Exception as e:
            raise ConfigurationError(f"Error cargando configuración: {e}")
        
        # Validar CV base
        if not self.validar_cv_base():
            raise FileProcessingError("CV base no válido")
            
        os.makedirs(self.carpeta_salida, exist_ok=True)
        
        # Inicializar base de datos
        self.db_path = "aplicaciones.db"
        self.inicializar_base_datos()
        
        # Adaptaciones del CV según el tipo de posición
        self.adaptaciones_cv = {
            'qa_automatizacion': {
                'titulo': 'QA Automation Engineer & Full Stack Developer',
                'resumen_adicional': 'Especializado en automatización de pruebas con Selenium, desarrollo de frameworks de testing y desarrollo full stack. Experiencia combinando desarrollo y QA en proyectos críticos.',
                'experiencias_extra': [
                    'Desarrollo de suites automatizadas usando Selenium WebDriver con Java y Python',
                    'Implementación de Page Object Model para mantener código reutilizable y escalable',
                    'Automatización de pruebas de APIs REST con validaciones de respuesta y datos',
                    'Desarrollo y testing simultáneo de funcionalidades en proyectos municipales'
                ]
            },
            'qa_manual': {
                'titulo': 'QA Manual Engineer & Full Stack Developer',
                'resumen_adicional': 'Especializado en testing funcional frontend y backend con enfoque en validación de datos críticos. Experiencia única combinando desarrollo y testing manual.',
                'experiencias_extra': [
                    'Ejecución de casos de prueba manuales en sistemas críticos desarrollados internamente',
                    'Validaciones cruzadas entre sistemas y bases de datos con conocimiento profundo del código',
                    'Documentación detallada de defectos y evidencias funcionales',
                    'Testing de funcionalidades desarrolladas en Python/FastAPI y Next.js'
                ]
            },
            'desarrollador_python': {
                'titulo': 'Python Full Stack Developer & QA Engineer',
                'resumen_adicional': 'Desarrollador especializado en Python con enfoque en calidad. Experiencia en FastAPI, desarrollo de APIs REST e integración con bases de datos PostgreSQL, complementado con expertise en testing.',
                'experiencias_extra': [
                    'Desarrollo de APIs REST escalables con Python y FastAPI en proyectos municipales',
                    'Implementación de microservicios con testing integrado desde el desarrollo',
                    'Optimización de consultas SQL y gestión de bases de datos PostgreSQL',
                    'Desarrollo de dashboards y paneles administrativos con Vue.js y Quasar'
                ]
            },
            'desarrollador_java': {
                'titulo': 'Java Developer & QA Engineer',  # Se ajustará según el nivel
                'resumen_adicional': 'Desarrollador con base sólida en Python/FastAPI y experiencia en QA. Mi experiencia en desarrollo backend y metodologías de testing me proporciona una excelente base para trabajar con Java.',
                'experiencias_extra': [
                    'Experiencia transferible desde Python/FastAPI hacia ecosistema Java/Spring',
                    'Base sólida en desarrollo de APIs REST y microservicios',
                    'Experiencia práctica en testing que complementa el desarrollo',
                    'Conocimiento en metodologías ágiles y mejores prácticas de desarrollo'
                ]
            }
        }

    def cargar_configuracion(self, config_path: str) -> Dict[str, Any]:
        """Carga la configuración desde archivo JSON y procesa variables de entorno"""
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                config_content = f.read()
            
            # Procesar variables de entorno
            config_content = self.procesar_variables_entorno(config_content)
            
            # Parsear JSON
            config = json.loads(config_content)
            
            return config
        except FileNotFoundError:
            raise ConfigurationError(f"Archivo de configuración no encontrado: {config_path}")
        except json.JSONDecodeError as e:
            raise ConfigurationError(f"Error en formato JSON: {e}")

    def procesar_variables_entorno(self, content: str) -> str:
        """Reemplaza variables de entorno en el contenido del config"""
        import re
        
        # Buscar patrones ${VARIABLE}
        def replace_env_var(match):
            var_name = match.group(1)
            env_value = os.getenv(var_name, match.group(0))  # Si no existe, mantener original
            
            # Convertir tipos específicos
            if var_name in ['EMAIL_ENABLED']:
                return 'true' if env_value.lower() in ['true', '1', 'yes', 'on'] else 'false'
            elif var_name in ['SMTP_PORT', 'UMBRAL_FIT_DEFAULT']:
                try:
                    return str(int(env_value))
                except (ValueError, TypeError):
                    return env_value
            else:
                return env_value
        
        # Reemplazar todas las variables ${VAR_NAME}
        processed_content = re.sub(r'\$\{([^}]+)\}', replace_env_var, content)
        
        return processed_content

    def validar_cv_base(self) -> bool:
        """Valida que el CV base exista y sea válido"""
        try:
            if not os.path.exists(self.cv_base_path):
                raise FileProcessingError(f"CV base no encontrado: {self.cv_base_path}")
            
            # Verificar que sea un archivo Word válido
            try:
                doc = Document(self.cv_base_path)
                if len(doc.paragraphs) < 5:
                    raise FileProcessingError("CV base parece estar vacío o corrupto")
            except Exception as e:
                raise FileProcessingError(f"Error leyendo CV base: {e}")
            
            logging.info(f"CV base validado: {self.cv_base_path}")
            return True
            
        except FileProcessingError as e:
            logging.error(f"Error validando CV base: {e}")
            return False

    def detectar_salario(self, texto_postulacion: str) -> Dict[str, Any]:
        """Detecta rangos salariales en la postulación"""
        resultado = {
            'salario_detectado': False,
            'moneda': None,
            'rango_min': None,
            'rango_max': None,
            'es_competitivo': None,
            'alertas': []
        }
        
        texto = texto_postulacion.lower()
        
        # Patrones para detectar salarios
        patrones_usd = [
            r'usd\s*(\d+(?:\.\d{3})*)',
            r'dolares?\s*(\d+(?:\.\d{3})*)',
            r'\$\s*(\d+(?:\.\d{3})*)\s*usd'
        ]
        
        patrones_ars = [
            r'\$\s*(\d+(?:\.\d{3})*)',
            r'pesos\s*(\d+(?:\.\d{3})*)',
            r'ars\s*(\d+(?:\.\d{3})*)'
        ]
        
        # Buscar USD primero
        for patron in patrones_usd:
            matches = re.findall(patron, texto)
            if matches:
                resultado['salario_detectado'] = True
                resultado['moneda'] = 'USD'
                salarios = [int(m.replace('.', '')) for m in matches]
                resultado['rango_min'] = min(salarios)
                resultado['rango_max'] = max(salarios) if len(salarios) > 1 else None
                break
        
        # Si no encontró USD, buscar ARS
        if not resultado['salario_detectado']:
            for patron in patrones_ars:
                matches = re.findall(patron, texto)
                if matches:
                    resultado['salario_detectado'] = True
                    resultado['moneda'] = 'ARS'
                    salarios = [int(m.replace('.', '')) for m in matches]
                    resultado['rango_min'] = min(salarios)
                    resultado['rango_max'] = max(salarios) if len(salarios) > 1 else None
                    break
        
        # Evaluar competitividad
        if resultado['salario_detectado'] and resultado['moneda'] == 'USD':
            min_esperado = self.config['deteccion_salarios']['salario_minimo_esperado_usd']
            max_esperado = self.config['deteccion_salarios']['salario_maximo_esperado_usd']
            
            if resultado['rango_min'] < min_esperado:
                resultado['alertas'].append(f"💰 Salario bajo: ${resultado['rango_min']} USD (mínimo esperado: ${min_esperado})")
                resultado['es_competitivo'] = False
            elif resultado['rango_min'] > max_esperado:
                resultado['alertas'].append(f"🎯 Salario alto: ${resultado['rango_min']} USD (máximo esperado: ${max_esperado})")
                resultado['es_competitivo'] = True
            else:
                resultado['es_competitivo'] = True
                resultado['alertas'].append(f"✅ Salario competitivo: ${resultado['rango_min']} USD")
        
        return resultado

    def inicializar_base_datos(self):
        """Inicializa la base de datos SQLite"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Crear tabla de aplicaciones
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS aplicaciones (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    empresa TEXT NOT NULL,
                    tipo_posicion TEXT NOT NULL,
                    nivel_seniority TEXT NOT NULL,
                    fecha_aplicacion DATETIME NOT NULL,
                    fit_percentage INTEGER NOT NULL,
                    salario_detectado REAL,
                    moneda TEXT,
                    keywords TEXT,
                    cv_path TEXT,
                    postulacion_path TEXT,
                    estado TEXT DEFAULT 'enviado',
                    notas TEXT,
                    fecha_respuesta DATETIME,
                    fecha_entrevista DATETIME,
                    resultado_final TEXT
                )
            ''')
            
            # Crear tabla de estadísticas
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS estadisticas_diarias (
                    fecha DATE PRIMARY KEY,
                    aplicaciones_enviadas INTEGER DEFAULT 0,
                    entrevistas_obtenidas INTEGER DEFAULT 0,
                    ofertas_recibidas INTEGER DEFAULT 0,
                    fit_promedio REAL DEFAULT 0
                )
            ''')
            
            conn.commit()
            conn.close()
            logging.info("Base de datos inicializada correctamente")
            
        except Exception as e:
            logging.error(f"Error inicializando base de datos: {e}")
            raise FileProcessingError(f"Error con base de datos: {e}")

    def guardar_aplicacion_db(self, empresa: str, tipo_posicion: str, nivel: str, 
                            fit_percentage: int, salario_info: Dict, keywords: List[str],
                            cv_path: str, postulacion_path: str) -> int:
        """Guarda una aplicación en la base de datos"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Convertir datetime a string para evitar warnings
            fecha_actual = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            fecha_hoy_str = datetime.now().strftime("%Y-%m-%d")
            
            cursor.execute('''
                INSERT INTO aplicaciones (
                    empresa, tipo_posicion, nivel_seniority, fecha_aplicacion,
                    fit_percentage, salario_detectado, moneda, keywords,
                    cv_path, postulacion_path
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                empresa, tipo_posicion, nivel, fecha_actual,
                fit_percentage, 
                salario_info.get('rango_min'), 
                salario_info.get('moneda'),
                ', '.join(keywords),
                cv_path, postulacion_path
            ))
            
            aplicacion_id = cursor.lastrowid
            
            # Actualizar estadísticas diarias
            cursor.execute('''
                INSERT OR IGNORE INTO estadisticas_diarias (fecha, aplicaciones_enviadas, fit_promedio)
                VALUES (?, 1, ?)
            ''', (fecha_hoy_str, fit_percentage))
            
            cursor.execute('''
                UPDATE estadisticas_diarias 
                SET aplicaciones_enviadas = aplicaciones_enviadas + 1,
                    fit_promedio = (
                        SELECT AVG(fit_percentage) 
                        FROM aplicaciones 
                        WHERE DATE(fecha_aplicacion) = ?
                    )
                WHERE fecha = ?
            ''', (fecha_hoy_str, fecha_hoy_str))
            
            conn.commit()
            conn.close()
            
            logging.info(f"Aplicacion guardada en DB: ID {aplicacion_id}")
            return aplicacion_id
            
        except Exception as e:
            logging.error(f"Error guardando en DB: {e}")
            return 0

    def obtener_estadisticas(self) -> Dict[str, Any]:
        """Obtiene estadísticas generales de aplicaciones"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Estadísticas generales
            cursor.execute('SELECT COUNT(*) FROM aplicaciones')
            total_aplicaciones = cursor.fetchone()[0]
            
            cursor.execute('SELECT AVG(fit_percentage) FROM aplicaciones')
            fit_promedio = cursor.fetchone()[0] or 0
            
            # Por tipo de posición
            cursor.execute('''
                SELECT tipo_posicion, COUNT(*), AVG(fit_percentage), AVG(salario_detectado)
                FROM aplicaciones 
                GROUP BY tipo_posicion
                ORDER BY COUNT(*) DESC
            ''')
            por_tipo = cursor.fetchall()
            
            # Por empresa
            cursor.execute('''
                SELECT empresa, COUNT(*), MAX(fecha_aplicacion)
                FROM aplicaciones 
                GROUP BY empresa
                ORDER BY COUNT(*) DESC
                LIMIT 10
            ''')
            por_empresa = cursor.fetchall()
            
            # Últimas 7 aplicaciones
            cursor.execute('''
                SELECT empresa, tipo_posicion, fit_percentage, fecha_aplicacion
                FROM aplicaciones 
                ORDER BY fecha_aplicacion DESC
                LIMIT 7
            ''')
            ultimas_aplicaciones = cursor.fetchall()
            
            # Estadísticas salariales
            cursor.execute('''
                SELECT moneda, AVG(salario_detectado), MIN(salario_detectado), MAX(salario_detectado)
                FROM aplicaciones 
                WHERE salario_detectado IS NOT NULL
                GROUP BY moneda
            ''')
            estadisticas_salarios = cursor.fetchall()
            
            conn.close()
            
            return {
                'total_aplicaciones': total_aplicaciones,
                'fit_promedio': round(fit_promedio, 1),
                'por_tipo_posicion': por_tipo,
                'por_empresa': por_empresa,
                'ultimas_aplicaciones': ultimas_aplicaciones,
                'estadisticas_salarios': estadisticas_salarios
            }
            
        except Exception as e:
            logging.error(f"Error obteniendo estadísticas: {e}")
            return {}

    def mostrar_dashboard(self):
        """Muestra dashboard de estadísticas en consola"""
        stats = self.obtener_estadisticas()
        
        if not stats:
            print("❌ No hay datos para mostrar")
            return
            
        print("\n" + "="*60)
        print("📊 DASHBOARD DE ESTADÍSTICAS")
        print("="*60)
        
        # Resumen general
        print(f"\n📈 RESUMEN GENERAL:")
        print(f"   • Total aplicaciones: {stats['total_aplicaciones']}")
        print(f"   • Fit promedio: {stats['fit_promedio']}%")
        
        # Por tipo de posición
        if stats['por_tipo_posicion']:
            print(f"\n🎯 POR TIPO DE POSICIÓN:")
            for tipo, count, fit_avg, salario_avg in stats['por_tipo_posicion']:
                salario_str = f" | Salario prom: ${int(salario_avg)}" if salario_avg else ""
                print(f"   • {tipo}: {count} aplicaciones | Fit: {fit_avg:.1f}%{salario_str}")
        
        # Por empresa
        if stats['por_empresa']:
            print(f"\n🏢 TOP EMPRESAS:")
            for empresa, count, ultima_fecha in stats['por_empresa'][:5]:
                print(f"   • {empresa}: {count} aplicaciones | Última: {ultima_fecha[:10]}")
        
        # Estadísticas salariales
        if stats['estadisticas_salarios']:
            print(f"\n💰 SALARIOS DETECTADOS:")
            for moneda, promedio, minimo, maximo in stats['estadisticas_salarios']:
                print(f"   • {moneda}: ${int(promedio)} promedio (${int(minimo)} - ${int(maximo)})")
        
        # Últimas aplicaciones
        if stats['ultimas_aplicaciones']:
            print(f"\n🕒 ÚLTIMAS APLICACIONES:")
            for empresa, tipo, fit, fecha in stats['ultimas_aplicaciones']:
                print(f"   • {empresa} ({tipo}) | {fit}% | {fecha[:10]}")
        
        print("\n" + "="*60)

    def procesar_batch_csv(self, archivo_csv: str) -> Dict[str, Any]:
        """Procesa múltiples postulaciones desde archivo CSV"""
        if not os.path.exists(archivo_csv):
            raise FileProcessingError(f"Archivo CSV no encontrado: {archivo_csv}")
        
        resultados = {
            'procesadas': 0,
            'exitosas': 0,
            'rechazadas': 0,
            'errores': 0,
            'detalle': []
        }
        
        try:
            with open(archivo_csv, 'r', encoding='utf-8') as file:
                reader = csv.DictReader(file)
                
                for row in reader:
                    empresa = row.get('empresa', '').strip()
                    descripcion = row.get('descripcion', '').strip()
                    
                    if not empresa or not descripcion:
                        print(f"⚠️ Fila incompleta ignorada: {empresa}")
                        continue
                    
                    resultados['procesadas'] += 1
                    print(f"\n{'='*50}")
                    print(f"📝 Procesando {resultados['procesadas']}: {empresa}")
                    
                    try:
                        resultado = self.procesar_postulacion(descripcion, empresa)
                        
                        if resultado:
                            resultados['exitosas'] += 1
                            resultados['detalle'].append({
                                'empresa': empresa,
                                'estado': 'exitosa',
                                'tipo_posicion': resultado['tipo_posicion'],
                                'cv_path': resultado['cv_path']
                            })
                            print(f"✅ {empresa}: CV generado exitosamente")
                        else:
                            resultados['rechazadas'] += 1
                            resultados['detalle'].append({
                                'empresa': empresa,
                                'estado': 'rechazada',
                                'razon': 'Fit insuficiente o fuera de estrategia'
                            })
                            print(f"❌ {empresa}: Rechazada (fit insuficiente)")
                            
                    except Exception as e:
                        resultados['errores'] += 1
                        resultados['detalle'].append({
                            'empresa': empresa,
                            'estado': 'error',
                            'razon': str(e)
                        })
                        print(f"💥 {empresa}: Error - {e}")
                        logging.error(f"Error procesando {empresa}: {e}")
        
        except Exception as e:
            raise FileProcessingError(f"Error leyendo CSV: {e}")
        
        return resultados

    def mostrar_resumen_batch(self, resultados: Dict[str, Any]):
        """Muestra resumen de procesamiento batch"""
        print("\n" + "="*60)
        print("📊 RESUMEN DE PROCESAMIENTO BATCH")
        print("="*60)
        
        print(f"\n📈 ESTADÍSTICAS:")
        print(f"   • Total procesadas: {resultados['procesadas']}")
        print(f"   • ✅ Exitosas: {resultados['exitosas']}")
        print(f"   • ❌ Rechazadas: {resultados['rechazadas']}")
        print(f"   • 💥 Errores: {resultados['errores']}")
        
        if resultados['procesadas'] > 0:
            tasa_exito = (resultados['exitosas'] / resultados['procesadas']) * 100
            print(f"   • 📊 Tasa de éxito: {tasa_exito:.1f}%")
        
        # Detalles por estado
        if resultados['exitosas'] > 0:
            print(f"\n✅ APLICACIONES EXITOSAS:")
            for detalle in resultados['detalle']:
                if detalle['estado'] == 'exitosa':
                    print(f"   • {detalle['empresa']} ({detalle['tipo_posicion']})")
        
        if resultados['rechazadas'] > 0:
            print(f"\n❌ APLICACIONES RECHAZADAS:")
            for detalle in resultados['detalle']:
                if detalle['estado'] == 'rechazada':
                    print(f"   • {detalle['empresa']}: {detalle['razon']}")
        
        if resultados['errores'] > 0:
            print(f"\n💥 ERRORES:")
            for detalle in resultados['detalle']:
                if detalle['estado'] == 'error':
                    print(f"   • {detalle['empresa']}: {detalle['razon']}")
        
        print("\n" + "="*60)

    def enviar_email_aplicacion(self, empresa: str, posicion: str, cv_path: str, speech: str, 
                               email_destino: str = None) -> bool:
        """Envía email de aplicación con CV adjunto"""
        
        if not self.config['email_config']['enabled']:
            print("📧 Email deshabilitado en configuración")
            return False
        
        if not email_destino:
            email_destino = input(f"📧 Email para {empresa} (enter para omitir): ").strip()
            if not email_destino:
                print("⏭️ Envío de email omitido")
                return False
        
        try:
            email_config = self.config['email_config']
            
            # Crear mensaje
            msg = MIMEMultipart()
            msg['From'] = email_config['email']
            msg['To'] = email_destino
            
            # Preparar variables para el template
            tecnologias = self.extraer_tecnologias_principales(posicion)
            variables = {
                'posicion': posicion,
                'empresa': empresa,
                'nombre_completo': email_config['nombre_completo'],
                'telefono': email_config['telefono'],
                'email': email_config['email'],
                'tecnologias_principales': ', '.join(tecnologias),
                'speech_personalizado': speech
            }
            
            # Asunto y cuerpo
            asunto = email_config['templates']['asunto'].format(**variables)
            cuerpo = email_config['templates']['cuerpo'].format(**variables)
            
            msg['Subject'] = asunto
            msg.attach(MIMEText(cuerpo, 'plain', 'utf-8'))
            
            # Adjuntar CV
            if os.path.exists(cv_path):
                with open(cv_path, "rb") as attachment:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(attachment.read())
                
                encoders.encode_base64(part)
                filename = os.path.basename(cv_path)
                part.add_header(
                    'Content-Disposition',
                    f'attachment; filename= {filename}'
                )
                msg.attach(part)
            
            # Enviar email
            server = smtplib.SMTP(email_config['smtp_server'], email_config['smtp_port'])
            server.starttls()
            server.login(email_config['email'], email_config['password'])
            text = msg.as_string()
            server.sendmail(email_config['email'], email_destino, text)
            server.quit()
            
            print(f"📧 ✅ Email enviado exitosamente a {email_destino}")
            logging.info(f"Email enviado a {empresa}: {email_destino}")
            return True
            
        except Exception as e:
            print(f"📧 ❌ Error enviando email: {e}")
            logging.error(f"Error enviando email a {empresa}: {e}")
            return False

    def extraer_tecnologias_principales(self, posicion: str) -> List[str]:
        """Extrae las tecnologías principales mencionadas en la posición"""
        tecnologias_destacadas = []
        posicion_lower = posicion.lower()
        
        # Buscar tecnologías clave en la posición
        tech_importantes = {
            'Python': ['python', 'fastapi', 'django', 'flask'],
            'Java': ['java', 'spring', 'spring boot'],
            'JavaScript': ['javascript', 'vue', 'react', 'angular', 'next.js'],
            'QA': ['qa', 'testing', 'selenium', 'automation'],
            'SQL': ['sql', 'postgresql', 'mysql', 'database']
        }
        
        for tech_name, keywords in tech_importantes.items():
            if any(kw in posicion_lower for kw in keywords):
                tecnologias_destacadas.append(tech_name)
        
        return tecnologias_destacadas[:3]  # Máximo 3 tecnologías principales

    def limpiar_nombre_archivo(self, texto: str) -> str:
        """Limpia texto para usarlo como nombre de archivo válido"""
        if not texto:
            return "sin_nombre"
        
        # Caracteres problemáticos en Windows
        caracteres_invalidos = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
        
        # Reemplazar caracteres problemáticos
        texto_limpio = texto
        for char in caracteres_invalidos:
            texto_limpio = texto_limpio.replace(char, '_')
        
        # Limpiar espacios múltiples y caracteres especiales adicionales
        texto_limpio = re.sub(r'[^\w\s\-_\.]', '_', texto_limpio)  # Solo alfanuméricos, espacios, guiones, puntos
        texto_limpio = re.sub(r'\s+', '_', texto_limpio)  # Espacios a guiones bajos
        texto_limpio = re.sub(r'_+', '_', texto_limpio)  # Múltiples _ a uno solo
        texto_limpio = texto_limpio.strip('_')  # Remover _ al inicio/final
        
        # Limitar longitud (nombres de archivo muy largos pueden fallar)
        if len(texto_limpio) > 50:
            texto_limpio = texto_limpio[:50].rstrip('_')
        
        # Fallback si queda vacío
        if not texto_limpio:
            texto_limpio = "empresa"
            
        return texto_limpio.lower()

    def obtener_headers_aleatorios(self) -> Dict[str, str]:
        """Obtiene headers aleatorios para evitar detección"""
        user_agents = [
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:89.0) Gecko/20100101 Firefox/89.0',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        ]
        
        import random
        return {
            'User-Agent': random.choice(user_agents),
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'es-ES,es;q=0.9,en;q=0.8',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
            'Sec-Fetch-Dest': 'document',
            'Sec-Fetch-Mode': 'navigate',
            'Sec-Fetch-Site': 'none',
            'Cache-Control': 'max-age=0'
        }

    def scrape_portal(self, portal_name: str, query: str, location: str = "Buenos Aires") -> List[Dict[str, str]]:
        """Scraping de un portal específico de trabajo"""
        if not self.config['scraping_config']['enabled']:
            print(f"🕷️ Web scraping deshabilitado en configuración")
            return []
            
        portal_config = self.config['scraping_config']['portales'].get(portal_name)
        if not portal_config or not portal_config['enabled']:
            print(f"⚠️ Portal {portal_name} no está habilitado")
            return []
        
        jobs = []
        headers = self.obtener_headers_aleatorios()
        
        try:
            # Limpiar query para URLs (especialmente ZoneJobs que no acepta espacios)
            query_clean = query.lower().replace(' ', '-').replace('_', '-')
            
            # Construir URL de búsqueda
            if portal_name == 'zonajobs':
                # ZoneJobs usa formato especial: empleos-busqueda-qa.html
                search_url = portal_config['search_url'].format(query=query_clean)
            else:
                # Otros portales usan query normal con URL encoding
                search_url = portal_config['search_url'].format(
                    query=requests.utils.quote(query),
                    location=requests.utils.quote(location)
                )
            
            print(f"🕷️ Scrapeando {portal_name}: {query} en {location}")
            print(f"   📍 URL: {search_url}")
            logging.info(f"Scraping {portal_name}: {search_url}")
            
            # Realizar request con mejor manejo de errores
            response = requests.get(search_url, headers=headers, timeout=15)
            
            # Verificar respuesta
            if response.status_code == 403:
                print(f"   ❌ 403 Forbidden - {portal_name} bloquea scraping")
                logging.warning(f"{portal_name} bloquea scraping: 403 Forbidden")
                return []
            elif response.status_code == 404:
                print(f"   ❌ 404 Not Found - URL incorrecta para {portal_name}")
                logging.warning(f"{portal_name} URL incorrecta: 404")
                return []
            elif response.status_code != 200:
                print(f"   ❌ Error {response.status_code} - {portal_name}")
                logging.warning(f"{portal_name} error HTTP: {response.status_code}")
                return []
            
            # Verificar que el contenido no esté vacío
            if len(response.content) < 1000:
                print(f"   ⚠️ Respuesta muy pequeña de {portal_name} - posible problema")
                logging.warning(f"{portal_name} respuesta pequeña: {len(response.content)} bytes")
            
            response.raise_for_status()
            
            # Parsear HTML
            soup = BeautifulSoup(response.content, 'html.parser')
            selectors = portal_config['job_selectors']
            
            # Buscar contenedores de trabajos
            job_containers = soup.select(selectors['job_container'])
            max_results = self.config['scraping_config']['max_results_per_portal']
            
            # DEBUG: Mostrar información sobre el HTML recibido
            print(f"   🔍 HTML recibido: {len(response.content)} bytes")
            print(f"   🎯 Selector usado: '{selectors['job_container']}'")
            print(f"   📦 Contenedores encontrados: {len(job_containers)}")
            
            # Si no encuentra trabajos, hacer debug más detallado
            if len(job_containers) == 0:
                print(f"   🚨 DEBUG: No se encontraron contenedores con selector '{selectors['job_container']}'")
                
                # Mostrar algunos selectores comunes para debug
                common_selectors = ['.job', '.aviso', '.offer', '.resultado', '.listado', '.item', 
                                  '[data-job]', '.trabajo', '.empleo', '.card']
                
                for sel in common_selectors:
                    found = soup.select(sel)
                    if len(found) > 0:
                        print(f"   💡 Selector alternativo '{sel}': {len(found)} elementos")
                        
                # Mostrar parte del HTML para debug manual
                print(f"   📄 Primeros 500 chars del HTML:")
                print(f"   {str(soup)[:500]}...")
                
                return []
            
            for i, container in enumerate(job_containers[:max_results]):
                try:
                    # Extraer información del trabajo
                    title_elem = container.select_one(selectors['title'])
                    company_elem = container.select_one(selectors['company']) 
                    desc_elem = container.select_one(selectors['description'])
                    salary_elem = container.select_one(selectors.get('salary', ''))
                    location_elem = container.select_one(selectors.get('location', ''))
                    
                    # Limpiar y extraer texto
                    title = title_elem.get_text(strip=True) if title_elem else "Sin título"
                    company = company_elem.get_text(strip=True) if company_elem else "Empresa confidencial"
                    description = desc_elem.get_text(strip=True) if desc_elem else ""
                    salary = salary_elem.get_text(strip=True) if salary_elem else ""
                    job_location = location_elem.get_text(strip=True) if location_elem else location
                    
                    # Limpiar datos
                    title = self.limpiar_texto(title)
                    company = self.limpiar_texto(company)
                    description = self.limpiar_texto(description)[:500]  # Limitar descripción
                    
                    # Filtrar trabajos spam
                    if self.es_trabajo_spam(title, company, description):
                        continue
                    
                    job_data = {
                        'portal': portal_name,
                        'title': title,
                        'company': company,
                        'description': description,
                        'salary': salary,
                        'location': job_location,
                        'url': search_url,
                        'scraped_at': datetime.now().isoformat()
                    }
                    
                    jobs.append(job_data)
                    
                except Exception as e:
                    logging.warning(f"Error procesando trabajo {i} de {portal_name}: {e}")
                    continue
            
            print(f"✅ {portal_name}: {len(jobs)} trabajos encontrados")
            
            # Delay entre requests para ser respetuosos
            delay = self.config['scraping_config']['delay_between_requests']
            if delay > 0:
                time.sleep(delay)
                
        except requests.exceptions.RequestException as e:
            print(f"❌ Error de red scrapeando {portal_name}: {e}")
            logging.error(f"Error de red en {portal_name}: {e}")
        except Exception as e:
            print(f"❌ Error inesperado scrapeando {portal_name}: {e}")
            logging.error(f"Error inesperado en {portal_name}: {e}")
        
        return jobs

    def limpiar_texto(self, texto: str) -> str:
        """Limpia texto extraído del scraping"""
        if not texto:
            return ""
        
        # Limpiar espacios y caracteres especiales
        texto = ' '.join(texto.split())
        texto = texto.replace('\n', ' ').replace('\r', '')
        texto = texto.replace('\t', ' ').strip()
        
        return texto[:200]  # Limitar longitud

    def es_trabajo_spam(self, title: str, company: str, description: str) -> bool:
        """Detecta trabajos spam o de baja calidad"""
        texto_completo = f"{title} {company} {description}".lower()
        
        # Palabras spam configurables
        palabras_spam = self.config['scraping_config']['filtros']['palabras_spam']
        
        for spam_word in palabras_spam:
            if spam_word.lower() in texto_completo:
                return True
        
        # Empresas a excluir
        empresas_excluir = self.config['scraping_config']['filtros']['excluir_empresas']
        for empresa_spam in empresas_excluir:
            if empresa_spam.lower() in company.lower():
                return True
        
        # Filtros adicionales básicos
        if len(title) < 5 or len(description) < 20:
            return True
            
        return False

    def buscar_trabajos_automatico(self, area_busqueda: str = "qa", ubicacion: str = "Buenos Aires") -> List[Dict[str, str]]:
        """Búsqueda automática en múltiples portales"""
        if not self.config['scraping_config']['enabled']:
            print("🕷️ Web scraping está deshabilitado")
            return []
        
        todos_trabajos = []
        keywords_area = self.config['scraping_config']['keywords_busqueda'].get(area_busqueda, [area_busqueda])
        
        print(f"\n🔍 BÚSQUEDA AUTOMÁTICA DE TRABAJOS")
        print(f"📍 Área: {area_busqueda}")
        print(f"🌎 Ubicación: {ubicacion}")
        print(f"🔑 Keywords: {', '.join(keywords_area)}")
        print("=" * 50)
        
        # Iterar por cada portal habilitado
        for portal_name, portal_config in self.config['scraping_config']['portales'].items():
            if not portal_config['enabled']:
                continue
                
            print(f"\n🕷️ Scrapeando {portal_name.upper()}...")
            
            # Buscar con cada keyword del área
            for keyword in keywords_area:
                try:
                    jobs = self.scrape_portal(portal_name, keyword, ubicacion)
                    todos_trabajos.extend(jobs)
                    
                    if jobs:
                        print(f"   └── '{keyword}': {len(jobs)} trabajos")
                    
                except Exception as e:
                    print(f"   └── ❌ Error con '{keyword}': {e}")
                    logging.error(f"Error buscando {keyword} en {portal_name}: {e}")
        
        # Eliminar duplicados basados en título + empresa
        trabajos_unicos = []
        seen = set()
        
        for trabajo in todos_trabajos:
            key = f"{trabajo['title']}_{trabajo['company']}".lower()
            if key not in seen:
                seen.add(key)
                trabajos_unicos.append(trabajo)
        
        print(f"\n📊 RESUMEN DE BÚSQUEDA:")
        print(f"   • Total encontrados: {len(todos_trabajos)}")
        print(f"   • Únicos (sin duplicados): {len(trabajos_unicos)}")
        print(f"   • Portales consultados: {len([p for p in self.config['scraping_config']['portales'] if self.config['scraping_config']['portales'][p]['enabled']])}")
        
        return trabajos_unicos

    def test_portales(self) -> Dict[str, bool]:
        """Testa la conectividad de todos los portales configurados"""
        print("\n🧪 TESTING DE PORTALES")
        print("="*50)
        
        resultados = {}
        headers = self.obtener_headers_aleatorios()
        
        for portal_name, portal_config in self.config['scraping_config']['portales'].items():
            if not portal_config['enabled']:
                print(f"⏭️ {portal_name}: Deshabilitado")
                resultados[portal_name] = False
                continue
            
            try:
                # Probar URL base
                base_url = portal_config['base_url']
                print(f"🌐 Testing {portal_name}: {base_url}")
                
                response = requests.get(base_url, headers=headers, timeout=10)
                
                if response.status_code == 200:
                    print(f"   ✅ Base OK ({response.status_code})")
                    
                    # Probar URL de búsqueda con query simple
                    if portal_name == 'zonajobs':
                        search_url = portal_config['search_url'].format(query='developer')
                    else:
                        search_url = portal_config['search_url'].format(
                            query='developer',
                            location='Buenos Aires'
                        )
                    
                    print(f"   🔍 Testing búsqueda: {search_url}")
                    search_response = requests.get(search_url, headers=headers, timeout=10)
                    
                    if search_response.status_code == 200:
                        print(f"   ✅ Búsqueda OK ({search_response.status_code}) - {len(search_response.content)} bytes")
                        resultados[portal_name] = True
                    else:
                        print(f"   ❌ Búsqueda FALLA ({search_response.status_code})")
                        resultados[portal_name] = False
                        
                else:
                    print(f"   ❌ Base FALLA ({response.status_code})")
                    resultados[portal_name] = False
                    
            except requests.exceptions.RequestException as e:
                print(f"   ❌ Error de conexión: {e}")
                resultados[portal_name] = False
            except Exception as e:
                print(f"   ❌ Error inesperado: {e}")
                resultados[portal_name] = False
            
            # Delay entre tests
            time.sleep(1)
        
        print(f"\n📊 RESUMEN DE TESTING:")
        working = sum(1 for v in resultados.values() if v)
        total = len(resultados)
        print(f"   • Portales funcionando: {working}/{total}")
        
        for portal, status in resultados.items():
            status_icon = "✅" if status else "❌"
            print(f"   • {portal}: {status_icon}")
        
        return resultados

    def debug_html_portal(self, portal_name: str, query: str = "qa") -> str:
        """Debug del HTML de un portal específico para encontrar selectores correctos"""
        portal_config = self.config['scraping_config']['portales'].get(portal_name)
        if not portal_config:
            print(f"❌ Portal {portal_name} no encontrado en configuración")
            return ""
        
        headers = self.obtener_headers_aleatorios()
        query_clean = query.lower().replace(' ', '-').replace('_', '-')
        
        try:
            # Construir URL
            if portal_name == 'zonajobs':
                search_url = portal_config['search_url'].format(query=query_clean)
            else:
                search_url = portal_config['search_url'].format(query=query, location="Buenos Aires")
            
            print(f"🔍 DEBUG HTML de {portal_name.upper()}")
            print(f"📍 URL: {search_url}")
            
            # Realizar request
            response = requests.get(search_url, headers=headers, timeout=15)
            print(f"📊 Status Code: {response.status_code}")
            print(f"📦 Tamaño: {len(response.content)} bytes")
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Guardar HTML completo para análisis
                debug_filename = f"debug_{portal_name}_{query}.html"
                with open(debug_filename, 'w', encoding='utf-8') as f:
                    f.write(soup.prettify())
                
                print(f"💾 HTML guardado en: {debug_filename}")
                
                # Buscar posibles selectores de trabajos
                print(f"\n🎯 BUSCANDO POSIBLES SELECTORES:")
                possible_selectors = [
                    # Selectores comunes de trabajos
                    '.job', '.aviso', '.offer', '.resultado', '.listado', '.item',
                    '.card', '.trabajo', '.empleo', '.position', '.vacancy',
                    '[data-job]', '[data-aviso]', '.search-result',
                    '.job-item', '.job-card', '.job-listing', '.job-offer',
                    
                    # Selectores específicos de ZoneJobs (basados en inspección manual)
                    '.avisoItem', '.resultado-aviso', '.aviso-resultado',
                    '.listado-item', '.trabajo-item', '.offer-item',
                    '.empleo-card', '.job-container', '.empleo-container',
                    
                    # Selectores por estructura HTML común
                    'article', 'section[class*="job"]', 'div[class*="aviso"]',
                    'div[class*="trabajo"]', 'div[class*="empleo"]', 'li[class*="job"]',
                    
                    # Selectores por ID o data attributes
                    '[id*="job"]', '[id*="aviso"]', '[data-*="job"]'
                ]
                
                for selector in possible_selectors:
                    elements = soup.select(selector)
                    if len(elements) > 0:
                        print(f"   💡 '{selector}': {len(elements)} elementos")
                        
                        # Mostrar contenido del primer elemento
                        if elements:
                            first_elem = elements[0]
                            text_preview = first_elem.get_text(strip=True)[:100]
                            print(f"      Ejemplo: {text_preview}...")
                
                # Análisis específico por portal
                print(f"\n🔧 ANÁLISIS ESPECÍFICO {portal_name.upper()}:")
                
                # Buscar títulos que contengan palabras relacionadas con trabajo
                job_titles = soup.find_all(string=lambda text: text and len(text.strip()) > 5 and any(
                    word in text.lower() for word in ['qa', 'tester', 'automation', 'engineer', 'developer', 
                                                     'programador', 'analista', 'junior', 'senior', 'ssr', 
                                                     'python', 'java', 'react', 'angular', 'vue']
                ))
                
                if job_titles:
                    print(f"   📋 Textos de trabajo encontrados: {len(job_titles)}")
                    for i, title in enumerate(job_titles[:5]):
                        clean_title = title.strip()
                        if len(clean_title) > 10:  # Solo títulos significativos
                            print(f"      {i+1}. {clean_title[:80]}...")
                            parent = title.parent
                            if parent:
                                classes = parent.get('class', [])
                                parent_info = f"{parent.name}"
                                if classes:
                                    parent_info += f" class='{' '.join(classes)}'"
                                print(f"         └── Padre: {parent_info}")
                                
                                # Buscar el contenedor abuelo
                                grandparent = parent.parent
                                if grandparent and grandparent.name not in ['body', 'html']:
                                    gp_classes = grandparent.get('class', [])
                                    gp_info = f"{grandparent.name}"
                                    if gp_classes:
                                        gp_info += f" class='{' '.join(gp_classes)}'"
                                    print(f"         └── Abuelo: {gp_info}")
                else:
                    print("   ❌ No se encontraron textos relacionados con trabajos")
                    
                # Buscar estructura principal del contenido
                print(f"\n🏗️ ANÁLISIS DE ESTRUCTURA HTML:")
                main_containers = soup.select('main, .main, #main, .content, .results, .listings, .jobs, .avisos')
                for container in main_containers[:3]:
                    if container:
                        classes = container.get('class', [])
                        id_val = container.get('id', '')
                        print(f"   📦 {container.name} id='{id_val}' class='{' '.join(classes)}'")
                        
                        # Buscar hijos directos que podrían ser trabajos
                        children = container.find_all(recursive=False)
                        print(f"      └── {len(children)} hijos directos")
                        for child in children[:3]:
                            child_classes = child.get('class', [])
                            child_id = child.get('id', '')
                            text_preview = child.get_text(strip=True)[:50]
                            print(f"         • {child.name} id='{child_id}' class='{' '.join(child_classes)}' text='{text_preview}...'")
                
                # Sugerir selectores más probables
                print(f"\n💡 SELECTORES MÁS PROBABLES para {portal_name.upper()}:")
                priority_selectors = []
                if portal_name == 'zonajobs':
                    priority_selectors = ['article', '.aviso', '.resultado', '.job', '.listado li', 'div[class*="aviso"]', 'section']
                elif portal_name == 'computrabajo':
                    priority_selectors = ['.box_offer', '.offer', '.job', '.resultado', 'article', '.aviso']
                elif portal_name == 'bumeran':
                    priority_selectors = ['.avisoItem', '.job-item', '.resultado', '.aviso', 'article', '.offer']
                
                for selector in priority_selectors:
                    found = soup.select(selector)
                    if found:
                        priority = "🔥 ALTA PRIORIDAD" if len(found) >= 3 and len(found) <= 20 else "⭐ Media prioridad"
                        print(f"   {priority} '{selector}': {len(found)} elementos")
                        
                        # Mostrar ejemplo del primer elemento
                        if found:
                            first = found[0]
                            text_sample = first.get_text(strip=True)[:60]
                            print(f"         Ejemplo: {text_sample}...")
                
                return debug_filename
            else:
                print(f"❌ Error HTTP: {response.status_code}")
                return ""
                
        except Exception as e:
            print(f"❌ Error en debug: {e}")
            return ""

    def guardar_trabajos_csv(self, trabajos: List[Dict[str, str]], filename: str = None) -> str:
        """Guarda trabajos encontrados en CSV para procesamiento batch"""
        if not trabajos:
            print("❌ No hay trabajos para guardar")
            return ""
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"trabajos_encontrados_{timestamp}.csv"
        
        filepath = os.path.join(self.carpeta_salida, filename)
        
        try:
            with open(filepath, 'w', newline='', encoding='utf-8') as csvfile:
                fieldnames = ['empresa', 'descripcion', 'portal', 'title', 'salary', 'location', 'url', 'scraped_at']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                
                writer.writeheader()
                for trabajo in trabajos:
                    # Adaptar formato para el procesador batch existente
                    writer.writerow({
                        'empresa': trabajo['company'],
                        'descripcion': f"{trabajo['title']} - {trabajo['description']}",
                        'portal': trabajo['portal'],
                        'title': trabajo['title'],
                        'salary': trabajo['salary'],
                        'location': trabajo['location'],
                        'url': trabajo['url'],
                        'scraped_at': trabajo['scraped_at']
                    })
            
            print(f"💾 Trabajos guardados en: {filepath}")
            logging.info(f"Trabajos guardados: {filepath}")
            return filepath
            
        except Exception as e:
            print(f"❌ Error guardando CSV: {e}")
            logging.error(f"Error guardando trabajos CSV: {e}")
            return ""

    def detectar_tipo_empresa(self, texto_postulacion: str, empresa: str) -> str:
        """Detecta el tipo de empresa basado en la postulación y nombre"""
        texto = (texto_postulacion + " " + empresa).lower()
        
        puntuaciones = {}
        
        for tipo_empresa, config in self.config['templates_empresa'].items():
            puntuacion = 0
            for keyword in config['keywords']:
                if keyword in texto:
                    puntuacion += 1
            puntuaciones[tipo_empresa] = puntuacion
        
        # Encontrar el tipo con mayor puntuación
        if max(puntuaciones.values()) > 0:
            tipo_detectado = max(puntuaciones, key=puntuaciones.get)
            logging.info(f"Tipo de empresa detectado: {tipo_detectado}")
            return tipo_detectado
        
        # Por defecto, si no detecta nada específico
        return 'producto'  # Más neutral

    def cargar_cv_base(self):
        """Carga el CV base desde archivo Word"""
        try:
            doc = Document(self.cv_base_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
        except Exception as e:
            print(f"Error cargando CV base: {e}")
            return ""

    def detectar_tipo_posicion(self, texto_postulacion):
        """Detecta el tipo de posición y nivel basado en el texto de la postulación"""
        texto = texto_postulacion.lower()
        
        # Contadores para cada tipo (solo áreas donde tenemos experiencia)
        puntos = {
            'qa_automatizacion': 0,
            'qa_manual': 0,
            'desarrollador_python': 0,
            'desarrollador_java': 0,
            'desarrollador_frontend': 0,
            'desarrollador_fullstack': 0
        }
        
        # Palabras clave para QA Automatización
        if any(kw in texto for kw in ['automatización', 'selenium', 'automatizador', 'automation', 'locust', 'cypress']):
            puntos['qa_automatizacion'] += 3
        
        # Palabras clave para QA Manual
        if any(kw in texto for kw in ['qa funcional', 'testing funcional', 'qa manual', 'casos de prueba']):
            puntos['qa_manual'] += 3
            
        # Primero verificar si hay tecnologías que NO conocemos
        tecnologias_no_conocidas = self.config['tecnologias_no_conocidas']
        
        if any(tech in texto for tech in tecnologias_no_conocidas):
            print(f">>> 🚫 Tecnologías detectadas fuera de nuestro perfil: {[tech for tech in tecnologias_no_conocidas if tech in texto]}")
            return None, None
        
        # Palabras clave para Python (más específicas)
        if any(kw in texto for kw in ['python', 'django', 'flask', 'fastapi', 'pandas', 'numpy']):
            puntos['desarrollador_python'] += 3
            
        # Palabras clave para Java
        if any(kw in texto for kw in ['java', 'spring', 'spring boot', 'hibernate']):
            puntos['desarrollador_java'] += 3
            
        # Palabras clave para Frontend
        if any(kw in texto for kw in ['vue.js', 'vue', 'angular', 'frontend', 'front-end', 'javascript', 'typescript']):
            puntos['desarrollador_frontend'] += 3
            
        # Palabras clave para Full Stack
        if any(kw in texto for kw in ['full stack', 'fullstack', 'full-stack']):
            puntos['desarrollador_fullstack'] += 3
        
        # Palabras generales QA
        if any(kw in texto for kw in ['qa', 'testing', 'pruebas', 'quality assurance']):
            puntos['qa_automatizacion'] += 1
            puntos['qa_manual'] += 1
            
        # Palabras generales desarrollo (solo si no hay tecnologías específicas detectadas)
        palabras_desarrollo_general = ['desarrollador', 'developer', 'programador']
        if any(kw in texto for kw in palabras_desarrollo_general):
            # Solo asignar puntos generales si ya hay alguna tecnología específica detectada
            if max(puntos.values()) > 0:
                for tipo in puntos:
                    if puntos[tipo] > 0:
                        puntos[tipo] += 1
            else:
                # Si no hay tecnologías específicas, no asignar puntos automáticamente
                print(">>> ⚠️ Menciona 'desarrollador' pero sin tecnologías específicas de nuestro perfil")
                pass
        
        # Si no detecta nada de nuestras categorías, rechazar automáticamente
        if max(puntos.values()) == 0:
            print(f">>> 🚫 POSICIÓN FUERA DE NUESTRAS ÁREAS DE EXPERIENCIA")
            print(f">>> Texto analizado: {texto[:200]}...")
            print(f">>> Solo aplicamos a: QA, Python, Java, Frontend, Full Stack")
            return None, None
        
        tipo_base = max(puntos, key=puntos.get)
        max_puntos = puntos[tipo_base]
        
        # Mostrar información de detección
        print(f">>> Detección: {tipo_base} (puntos: {max_puntos})")
        
        # Detectar nivel de seniority
        nivel = self.detectar_nivel_seniority(texto)
        
        return tipo_base, nivel
    
    def detectar_nivel_seniority(self, texto):
        """Detecta el nivel de seniority requerido"""
        texto_lower = texto.lower()
        
        # Primero detectar SSR y variantes específicas de semi-senior
        if any(kw in texto_lower for kw in ['ssr', 'semi senior', 'semi-senior', 'advance', 'intermedio']):
            return 'semi_senior'
        elif any(kw in texto_lower for kw in ['junior', 'jr', 'trainee', 'entry level', 'sin experiencia']):
            return 'junior'
        elif any(kw in texto_lower for kw in ['senior', ' sr ', 'lead', 'líder', 'tech lead']):
            return 'senior'
        else:
            return 'semi_senior'  # Default

    def extraer_keywords_avanzado(self, texto_postulacion):
        """Extrae keywords relevantes de la postulación"""
        texto = texto_postulacion.lower()
        keywords_encontradas = []
        
        # Buscar en todas las categorías
        for categoria, keywords in self.perfil_tecnico.items():
            for kw in keywords:
                if kw in texto:
                    keywords_encontradas.append(kw)
        
        # Si no encuentra nada, buscar palabras más generales
        if not keywords_encontradas:
            palabras_generales = ['qa', 'testing', 'pruebas', 'developer', 'desarrollador', 
                                'programador', 'java', 'python', 'sql', 'api', 'web', 
                                'frontend', 'backend', 'full stack', 'scrum', 'agile']
            
            for palabra in palabras_generales:
                if palabra in texto:
                    keywords_encontradas.append(palabra)
        
        if not keywords_encontradas:
            print(">>> ⚠️  No se detectaron keywords específicas - usando tipo por defecto")
        
        return list(set(keywords_encontradas))  # Eliminar duplicados
    
    def validar_estrategia_aplicacion(self, tipo_posicion, nivel):
        """Valida si la posición está dentro de la estrategia de aplicación"""
        
        estrategia = self.config['estrategia_aplicacion']
        
        if nivel not in estrategia:
            return False
            
        tipos_permitidos = estrategia[nivel]
        return tipo_posicion in tipos_permitidos
    
    def tiene_experiencia_frontend(self):
        """Verifica si tenemos experiencia suficiente en frontend para aplicar"""
        # Basado en tu experiencia: Vue.js, Angular, Next.js, Quasar
        return True  # Tienes experiencia real en frontend
    
    def generar_experiencias_tecnicas_especificas(self, keywords, tipo_posicion):
        """Genera experiencias técnicas específicas basadas en keywords encontradas"""
        experiencias = []
        
        # Git Flow y Code Review (muy demandado)
        if any(kw in keywords for kw in ['git', 'github', 'gitlab']) or 'desarrollador' in tipo_posicion:
            experiencias.append("Gestión de código con Git, incluyendo branching strategies y code review en proyectos colaborativos")
        
        # Deploy y CI/CD
        if any(kw in keywords for kw in ['deploy', 'ci/cd', 'jenkins', 'docker']):
            experiencias.append("Participación en procesos de deployment y gestión de ramas en entornos de desarrollo")
        
        # Java específico (para posiciones junior)
        if 'java' in keywords and tipo_posicion == 'desarrollador_java':
            experiencias.append("Base sólida en desarrollo backend transferible a Java/Spring Boot")
            experiencias.append("Experiencia en APIs REST aplicable al ecosistema Java")
        
        # Bases de datos específicas
        if any(kw in keywords for kw in ['postgresql', 'mysql', 'sql']):
            experiencias.append("Optimización de consultas SQL y diseño de esquemas de base de datos")
        
        # Metodologías específicas mencionadas
        if any(kw in keywords for kw in ['scrum', 'kanban', 'agile']):
            experiencias.append("Participación activa en ceremonias ágiles y trabajo colaborativo en equipos multidisciplinarios")
        
        # APIs y microservicios
        if any(kw in keywords for kw in ['api', 'rest', 'microservicios']):
            experiencias.append("Diseño e implementación de APIs REST siguiendo mejores prácticas de arquitectura")
        
        return experiencias
    
    def generar_analisis_fit(self, keywords_postulacion, tipo_posicion, nivel, empresa):
        """Genera análisis automático de fit entre CV y postulación"""
        
        # Keywords que tenemos en nuestro CV (basado en tu experiencia real)
        nuestras_fortalezas = ['python', 'fastapi', 'postgresql', 'vue', 'quasar', 'next.js', 
                              'locust', 'postman', 'qa', 'manual', 'testing', 'scrum', 'kanban', 'git',
                              'javascript', 'sql', 'api', 'rest', 'agile', 'full stack', 'desarrollador', 
                              'angular', 'webservices', 'selenium', 'automatización', 'automation',
                              'frontend', 'backend', 'casos de prueba', 'validaciones', 'metodologías']
        
        coincidencias = []
        brechas = []
        
        # Analizar coincidencias
        for kw in keywords_postulacion:
            if kw in nuestras_fortalezas:
                coincidencias.append(kw)
        
        # Detectar brechas específicas
        if 'java' in keywords_postulacion and tipo_posicion == 'desarrollador_java':
            if nivel == 'senior':
                brechas.append("Java Senior requiere más experiencia - experiencia principal en Python")
            elif nivel == 'semi_senior':
                # Para SSR, solo mencionar como área de crecimiento, no como brecha bloqueante
                if len(coincidencias) < 3:
                    brechas.append("Java SSR - aprovechar experiencia backend transferible desde Python")
        
        if any(kw in keywords_postulacion for kw in ['code review', 'git flow', 'deploy']):
            if not any(kw in nuestras_fortalezas for kw in ['git', 'github']):
                brechas.append("Falta experiencia explícita en Git Flow/Code Review")
        
        if nivel == 'senior' and len(coincidencias) < 3:
            brechas.append("Puede requerir más experiencia para nivel Senior")
        
        # Calcular porcentaje de fit mejorado
        if not keywords_postulacion:
            # Si no hay keywords detectadas, usar heurística basada en tipo de posición
            fit_percentage = 60  # Base mínima
            if tipo_posicion in ['qa_manual', 'qa_automatizacion']:
                fit_percentage = 75  # Tenemos experiencia sólida en QA
            elif tipo_posicion == 'desarrollador_python':
                fit_percentage = 85  # Nuestra fortaleza principal
        else:
            total_keywords = len(keywords_postulacion)
            coincidencias_count = len(coincidencias)
            
            # Para QA, usar lógica especial porque es nuestra fortaleza
            if tipo_posicion in ['qa_manual', 'qa_automatizacion']:
                # Base alta para QA porque es nuestra experiencia
                base_percentage = 65
                
                # Bonificaciones específicas para QA
                if any(kw in coincidencias for kw in ['qa', 'testing', 'automatización', 'selenium']):
                    base_percentage += 15
                if any(kw in coincidencias for kw in ['sql', 'api', 'rest']):
                    base_percentage += 10
                if any(kw in coincidencias for kw in ['git', 'scrum', 'agile']):
                    base_percentage += 5
                    
                # Si tiene Java pero es QA, no penalizar tanto
                if 'java' in keywords_postulacion and tipo_posicion == 'qa_automatizacion':
                    base_percentage += 5  # Bonus menor porque podemos aprender Java para QA
                    
            else:
                # Para desarrollo, usar cálculo mejorado
                base_percentage = (coincidencias_count / total_keywords * 100) if total_keywords > 0 else 0
                
                if tipo_posicion == 'desarrollador_python' and any(kw in keywords_postulacion for kw in ['python', 'fastapi']):
                    base_percentage += 15  # Nuestra tecnología principal
                elif tipo_posicion == 'desarrollador_java':
                    # Para Java, dar base más alta porque tenemos experiencia backend transferible
                    base_percentage = max(base_percentage, 50)  # Base mínima para Java
                    
                    # Bonificaciones para Java
                    if any(kw in coincidencias for kw in ['sql', 'base de datos']):
                        base_percentage += 10  # Tenemos experiencia SQL
                    if any(kw in coincidencias for kw in ['api', 'rest']):
                        base_percentage += 10  # Experiencia en APIs
                    if any(kw in coincidencias for kw in ['scrum', 'ágiles', 'agile']):
                        base_percentage += 5   # Metodologías ágiles
                    if 'java' in keywords_postulacion:
                        base_percentage += 5   # Bonus por mencionar Java específicamente
            
            fit_percentage = min(100, base_percentage)  # Cap a 100%
        
        return {
            'fit_percentage': round(fit_percentage),
            'coincidencias': coincidencias,
            'brechas': brechas,
            'recomendaciones': self.generar_recomendaciones(tipo_posicion, nivel, brechas)
        }
    
    def generar_recomendaciones(self, tipo_posicion, nivel, brechas):
        """Genera recomendaciones específicas para mejorar el fit"""
        recomendaciones = []
        
        if tipo_posicion == 'desarrollador_java':
            recomendaciones.append("Destacar experiencia transferible desde Python hacia Java/Spring")
            recomendaciones.append("Mencionar cursos o proyectos personales en Java si los hay")
        
        if 'Git Flow' in str(brechas):
            recomendaciones.append("Agregar experiencia en Git Flow y code review al CV")
        
        if nivel == 'junior':
            recomendaciones.append("Enfatizar disposición a aprender y adaptabilidad")
        elif nivel == 'senior':
            recomendaciones.append("Destacar liderazgo técnico y mentoría a otros desarrolladores")
        
        return recomendaciones

    def adaptar_cv(self, cv_base, tipo_posicion, nivel, keywords_encontradas, empresa, texto_postulacion):
        """Adapta el CV según el tipo de posición, nivel y tipo de empresa detectado"""
        adaptacion = self.adaptaciones_cv.get(tipo_posicion, self.adaptaciones_cv['qa_manual'])
        
        # Detectar tipo de empresa
        tipo_empresa = self.detectar_tipo_empresa(texto_postulacion, empresa)
        template_empresa = self.config['templates_empresa'][tipo_empresa]
        
        # 1. Modificar título principal según el nivel y tipo de empresa
        titulo_adaptado = adaptacion['titulo']
        if tipo_posicion == 'desarrollador_java':
            if nivel == 'junior':
                titulo_adaptado = 'Java Junior Developer & QA Engineer'
            elif nivel == 'semi_senior':
                titulo_adaptado = 'Java SSR Developer & QA Engineer'
            elif nivel == 'senior':
                titulo_adaptado = 'Java Senior Developer & QA Engineer'
        
        # Agregar sufijo según tipo de empresa
        titulo_adaptado += template_empresa['adaptaciones']['titulo_suffix']
        
        cv_adaptado = cv_base.replace('QA Engineer', titulo_adaptado)
        print(f">>> 🏢 Empresa tipo: {tipo_empresa} | Título adaptado: {titulo_adaptado}")
        
        # 2. Mejorar el perfil profesional integrando la especialización y tipo de empresa
        perfil_original = "QA Engineer y Full Stack Developer con experiencia en validación de datos, pruebas de sistemas y desarrollo de aplicaciones en entornos ágiles y arquitecturas de microservicios."
        
        perfil_mejorado = f"{adaptacion['titulo']} con experiencia en validación de datos, pruebas de sistemas y desarrollo de aplicaciones en entornos ágiles y arquitecturas de microservicios. {adaptacion['resumen_adicional']} {template_empresa['adaptaciones']['enfoque_experiencia']}"
        
        cv_adaptado = cv_adaptado.replace(perfil_original, perfil_mejorado)
        
        # 3. Agregar experiencias relevantes en la sección de logros
        logros_adicionales = "\n\nExperiencias Técnicas Destacadas:\n"
        for exp in adaptacion['experiencias_extra']:
            logros_adicionales += f"• {exp}\n"
        
        # Agregar experiencias específicas del tipo de empresa
        logros_adicionales += f"\nExperiencias orientadas a {tipo_empresa.title()}:\n"
        for exp_empresa in template_empresa['adaptaciones']['logros_adicionales']:
            logros_adicionales += f"• {exp_empresa}\n"
        
        # Agregar experiencias técnicas específicas según keywords
        logros_tecnicos = self.generar_experiencias_tecnicas_especificas(keywords_encontradas, tipo_posicion)
        for exp_tec in logros_tecnicos:
            logros_adicionales += f"• {exp_tec}\n"
        
        # Ajustar según nivel de seniority
        if nivel == 'junior':
            logros_adicionales += "• Enfoque en aprendizaje continuo y adaptación a nuevas tecnologías\n"
        elif nivel == 'senior':
            logros_adicionales += "• Mentoría a desarrolladores junior y liderazgo técnico en proyectos\n"
        
        # Insertar después de "Logros Relevantes"
        cv_adaptado = cv_adaptado.replace(
            "Desarrollo de funcionalidades completas en plataformas de compras municipales con stack completo Python/JavaScript.",
            f"Desarrollo de funcionalidades completas en plataformas de compras municipales con stack completo Python/JavaScript.{logros_adicionales}"
        )
        
        # 4. Agregar keywords relevantes sutilmente
        keywords_faltantes = []
        for kw in keywords_encontradas:
            if kw not in cv_adaptado.lower() and kw in ['selenium', 'java', 'spring boot', 'automatización', 'katalon', 'uft']:
                keywords_faltantes.append(kw)
        
        if keywords_faltantes:
            tech_adicional = f"\n\nTecnologías y herramientas relevantes que fui adquiriendo: {', '.join(keywords_faltantes).title()}"
            cv_adaptado += tech_adicional
        
        return cv_adaptado, adaptacion['titulo']

    def generar_cv_pdf(self, texto_cv, nombre_archivo):
        """Genera el CV en formato PDF con mejor formato"""
        try:
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
            from reportlab.lib import colors
            
            doc = SimpleDocTemplate(nombre_archivo, pagesize=A4, 
                                  leftMargin=50, rightMargin=50, 
                                  topMargin=50, bottomMargin=50)
            styles = getSampleStyleSheet()
            
            # Estilos personalizados
            titulo_style = ParagraphStyle(
                'TituloPersonal',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=6,
                alignment=TA_CENTER,
                textColor=colors.darkblue
            )
            
            subtitulo_style = ParagraphStyle(
                'SubtituloPersonal', 
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=8,
                textColor=colors.darkblue,
                borderWidth=1,
                borderColor=colors.lightgrey,
                borderPadding=3
            )
            
            contenido = []
            lineas = texto_cv.split('\n')
            
            for i, linea in enumerate(lineas):
                linea = linea.strip()
                if not linea:
                    contenido.append(Spacer(1, 6))
                    continue
                
                # Nombre (primera línea)
                if i == 0:
                    contenido.append(Paragraph(linea, titulo_style))
                    contenido.append(Spacer(1, 12))
                
                # Títulos de sección
                elif linea in ['Perfil Profesional', 'Logros Relevantes', 'Experiencia Profesional', 
                              'Tecnologías y Herramientas', 'Educación', 'Experiencias Técnicas Destacadas']:
                    contenido.append(Spacer(1, 12))
                    contenido.append(Paragraph(linea, subtitulo_style))
                    contenido.append(Spacer(1, 6))
                
                # Subtítulo de trabajo/empresa
                elif any(empresa in linea for empresa in ['Municipalidad', 'FABRICARG', 'CECAL', 'Proyecto Freelance']):
                    contenido.append(Spacer(1, 8))
                    contenido.append(Paragraph(f"<b>{linea}</b>", styles["Normal"]))
                    contenido.append(Spacer(1, 4))
                
                # Contenido normal
                else:
                    contenido.append(Paragraph(linea, styles["Normal"]))
                    contenido.append(Spacer(1, 4))
            
            doc.build(contenido)
            return True
        except Exception as e:
            print(f"Error generando PDF: {e}")
            return False

    def generar_speech_avanzado(self, empresa, tipo_posicion, nivel, keywords):
        """Genera un speech personalizado según el tipo de posición y nivel"""
        
        # Definir speech para Java según el nivel
        if tipo_posicion == 'desarrollador_java':
            if nivel == 'junior':
                speech_java = f"Gracias por la oportunidad en {empresa}. Me interesa esta posición junior porque mi experiencia sólida en Python/FastAPI y metodologías de testing me proporciona una excelente base para transicionar al ecosistema Java. Estoy motivado por aprender y aplicar mis conocimientos de backend en Java."
            else:  # semi_senior
                speech_java = f"Gracias por la oportunidad en {empresa}. Me motiva esta posición SSR porque mi experiencia en desarrollo backend con Python/FastAPI, combined con mi background en QA, me permite aportar una perspectiva integral al desarrollo Java. Mi experiencia en APIs REST y metodologías ágiles es directamente transferible al ecosistema Java/Spring."
        else:
            # Para casos que no sean Java, definir speech_java por defecto
            speech_java = f"Gracias por la oportunidad en {empresa}. Mi experiencia en desarrollo backend me proporciona una base sólida para trabajar con Java."
        
        speeches_base = {
            'qa_automatizacion': f"Gracias por la oportunidad en {empresa}. Me postulé porque tengo una experiencia única combinando desarrollo full stack y QA. Actualmente en la Municipalidad desarrollo funcionalidades con Python/FastAPI y Next.js, y también implemento automatización de pruebas con Locust y Selenium.",
            
            'qa_manual': f"Gracias por la oportunidad en {empresa}. Me entusiasma esta posición porque mi experiencia combina QA manual con desarrollo. En mis proyectos actuales desarrollo las funcionalidades y luego las testeo, lo que me da una perspectiva única para detectar problemas desde el diseño.",
            
            'desarrollador_python': f"Gracias por la oportunidad en {empresa}. Me motiva esta posición porque tengo experiencia sólida desarrollando con Python/FastAPI, PostgreSQL y frontend con Vue.js en proyectos municipales. Mi background en QA me permite desarrollar código más robusto desde el inicio.",
            
            'desarrollador_java': speech_java
        }
        
        speech_base = speeches_base.get(tipo_posicion, speeches_base['qa_manual'])
        
        speech_base += " Estoy entusiasmado por aportar valor al equipo y seguir creciendo profesionalmente en este rol."
        
        return speech_base

    def guardar_postulacion(self, texto_postulacion, empresa, tipo_posicion):
        """Guarda la postulación con metadatos"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        empresa_limpia = self.limpiar_nombre_archivo(empresa)
        nombre_archivo = f"postulacion_{empresa_limpia}_{tipo_posicion}_{timestamp}.txt"
        path_completo = os.path.join(self.carpeta_salida, nombre_archivo)
        
        contenido = f"""EMPRESA: {empresa}
TIPO POSICIÓN: {tipo_posicion}
FECHA: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
=====================================

{texto_postulacion}
"""
        
        with open(path_completo, "w", encoding="utf-8") as f:
            f.write(contenido)
        
        return path_completo

    def procesar_postulacion(self, texto_postulacion, empresa):
        """Proceso principal: analiza postulación y genera CV personalizado"""
        print(f"\n>>> Analizando postulación de {empresa}...")
        
        # 1. Detectar tipo de posición y nivel
        tipo_posicion, nivel = self.detectar_tipo_posicion(texto_postulacion)
        
        # Si no detectó una posición válida, terminar aquí
        if tipo_posicion is None:
            return None
            
        print(f">>> Tipo detectado: {tipo_posicion} ({nivel})")
        
        # 2. Extraer keywords
        keywords = self.extraer_keywords_avanzado(texto_postulacion)
        print(f">>> Keywords encontradas: {', '.join(keywords[:5])}{'...' if len(keywords) > 5 else ''}")
        
        # 2.5. Detectar salario
        try:
            info_salario = self.detectar_salario(texto_postulacion)
            if info_salario['salario_detectado']:
                print(f">>> Salario detectado: {info_salario['rango_min']} {info_salario['moneda']}")
                for alerta in info_salario['alertas']:
                    print(f">>> {alerta}")
            else:
                print(">>> No se detectó información salarial")
        except Exception as e:
            logging.warning(f"Error detectando salario: {e}")
        
        # 3. Cargar y adaptar CV
        cv_base = self.cargar_cv_base()
        if not cv_base:
            print(">>> Error: No se pudo cargar el CV base")
            return None
            
        # 4. Generar análisis de fit ANTES de crear archivos
        analisis_fit = self.generar_analisis_fit(keywords, tipo_posicion, nivel, empresa)
        print(f">>> Análisis de Fit: {analisis_fit['fit_percentage']}%")
        
        # 5. Validar estrategia de aplicación según nivel
        if not self.validar_estrategia_aplicacion(tipo_posicion, nivel):
            print(f"\n🚫 POSICIÓN FUERA DE ESTRATEGIA")
            print(f">>> Tipo detectado: {tipo_posicion} ({nivel})")
            print(">>> Estrategia actual:")
            print("   • Junior: QA, Python, Java, Frontend, Full Stack")
            print("   • Semi-Senior: QA, Python, Java, Full Stack")
            print("   • Rechazados automáticos: Senior y áreas sin experiencia")
            print(">>> No se generará CV para esta postulación")
            return None
        
        # 6. Validar umbral mínimo
        if analisis_fit['fit_percentage'] < self.umbral_fit:
            print(f"\n>>> FIT INSUFICIENTE ({analisis_fit['fit_percentage']}%)")
            print(f">>> Mínimo requerido: {self.umbral_fit}%")
            print(">>> No se generará CV para esta postulación")
            
            if analisis_fit['brechas']:
                print(f">>> Principales brechas: {', '.join(analisis_fit['brechas'])}")
            
            if analisis_fit['recomendaciones']:
                print("\n💡 Recomendaciones para mejorar fit:")
                for rec in analisis_fit['recomendaciones']:
                    print(f"   • {rec}")
            
            return None
        
        print(f"✅ FIT APROPIADO ({analisis_fit['fit_percentage']}%) - Generando CV...")
        
        # 6. Adaptar CV (solo si fit >= 70%)
        cv_adaptado, titulo_adaptado = self.adaptar_cv(cv_base, tipo_posicion, nivel, keywords, empresa, texto_postulacion)
        
        # 7. Generar archivos
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        empresa_limpia = self.limpiar_nombre_archivo(empresa)
        nombre_pdf = os.path.join(self.carpeta_salida, f"cv_{empresa_limpia}_{tipo_posicion}_{timestamp}.pdf")
        
        if self.generar_cv_pdf(cv_adaptado, nombre_pdf):
            print(f">>> CV generado: {nombre_pdf}")
        else:
            print(">>> Error generando PDF")
            return None
        
        # 8. Guardar postulación
        path_postulacion = self.guardar_postulacion(texto_postulacion, empresa, tipo_posicion)
        print(f">>> Postulación guardada: {path_postulacion}")
        
        # 9. Generar speech
        speech = self.generar_speech_avanzado(empresa, tipo_posicion, nivel, keywords)
        print(f"\n>>> Speech para entrevista:")
        print(f"'{speech}'\n")
        
        # Mostrar áreas de mejora si las hay
        if analisis_fit['brechas']:
            print(f">>> Áreas a considerar en entrevista: {', '.join(analisis_fit['brechas'])}")
        
        # 10. Guardar en base de datos
        try:
            self.guardar_aplicacion_db(
                empresa, tipo_posicion, nivel, analisis_fit['fit_percentage'],
                info_salario if 'info_salario' in locals() else {}, keywords,
                nombre_pdf, path_postulacion
            )
        except Exception as e:
            logging.warning(f"Error guardando en base de datos: {e}")
        
        # 11. Ofrecer envío de email
        if self.config['email_config']['enabled']:
            try:
                self.enviar_email_aplicacion(empresa, titulo_adaptado, nombre_pdf, speech)
            except Exception as e:
                logging.warning(f"Error en envío de email: {e}")
        
        # 12. Guardar resumen completo
        self.guardar_resumen(empresa, tipo_posicion, nivel, titulo_adaptado, keywords, speech, 
                           analisis_fit, nombre_pdf, path_postulacion)
        
        return {
            'empresa': empresa,
            'tipo_posicion': tipo_posicion,
            'titulo': titulo_adaptado,
            'keywords': keywords,
            'cv_path': nombre_pdf,
            'postulacion_path': path_postulacion,
            'speech': speech
        }

    def guardar_resumen(self, empresa, tipo_posicion, nivel, titulo, keywords, speech, analisis_fit, cv_path, postulacion_path):
        """Guarda un resumen completo de la postulación procesada"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        empresa_limpia = self.limpiar_nombre_archivo(empresa)
        resumen_path = os.path.join(self.carpeta_salida, f"resumen_{empresa_limpia}_{timestamp}.json")
        
        resumen = {
            'empresa': empresa,
            'fecha_procesamiento': datetime.now().isoformat(),
            'tipo_posicion': tipo_posicion,
            'nivel_seniority': nivel,
            'titulo_cv': titulo,
            'keywords_detectadas': keywords,
            'speech_entrevista': speech,
            'analisis_fit': analisis_fit,
            'archivos_generados': {
                'cv_pdf': cv_path,
                'postulacion': postulacion_path
            }
        }
        
        with open(resumen_path, 'w', encoding='utf-8') as f:
            json.dump(resumen, f, ensure_ascii=False, indent=2)

def parse_arguments():
    """Parsea argumentos de línea de comandos"""
    parser = argparse.ArgumentParser(
        description='Generador de CV Inteligente v3.0',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ejemplos de uso:
  python generador_cv_avanzado.py                          # Modo interactivo
  python generador_cv_avanzado.py --batch postulaciones.csv # Modo batch
  python generador_cv_avanzado.py --stats                   # Ver estadísticas
  python generador_cv_avanzado.py --scrape qa --save-jobs   # Buscar trabajos QA
  python generador_cv_avanzado.py --scrape python --location "Córdoba" # Python en Córdoba
  python generador_cv_avanzado.py --empresa "TechCorp" --postulacion "Descripción..." --email
  python generador_cv_avanzado.py --batch postulaciones.csv --email # Batch con emails
  python generador_cv_avanzado.py --umbral 80 --email       # Personalizar umbral y email
        """
    )
    
    parser.add_argument('--batch', '-b', 
                        help='Procesar múltiples postulaciones desde archivo CSV')
    parser.add_argument('--stats', '-s', action='store_true',
                        help='Mostrar dashboard de estadísticas')
    parser.add_argument('--empresa', '-e',
                        help='Nombre de la empresa (modo directo)')
    parser.add_argument('--postulacion', '-p',
                        help='Descripción de la postulación (modo directo)')
    parser.add_argument('--config', '-c', default='config.json',
                        help='Archivo de configuración (default: config.json)')
    parser.add_argument('--umbral', '-u', type=int,
                        help='Umbral mínimo de fit (override config)')
    parser.add_argument('--email', action='store_true',
                        help='Habilitar envío automático de emails')
    parser.add_argument('--email-to',
                        help='Email destino para modo directo')
    parser.add_argument('--scrape', '-w',
                        help='Buscar trabajos automáticamente (ej: qa, python, java)')
    parser.add_argument('--location', '-l', default='Buenos Aires',
                        help='Ubicación para búsqueda (default: Buenos Aires)')
    parser.add_argument('--save-jobs', action='store_true',
                        help='Guardar trabajos encontrados en CSV')
    parser.add_argument('--test-portales', action='store_true',
                        help='Testear conectividad de todos los portales')
    parser.add_argument('--debug-html', 
                        help='Debug HTML de un portal específico (ej: zonajobs)')
    
    return parser.parse_args()

# FUNCIÓN PRINCIPAL
def main():
    args = parse_arguments()
    
    try:
        # Inicializar generador con configuración específica
        generador = GeneradorCVInteligente(args.config)
        
        # Override configuraciones CLI
        if args.umbral:
            generador.umbral_fit = args.umbral
            print(f"🎯 Umbral de fit personalizado: {args.umbral}%")
        
        if args.email:
            generador.config['email_config']['enabled'] = True
            print(f"📧 Email habilitado por CLI")
            
        logging.info("Generador de CV iniciado correctamente")
    except (ConfigurationError, FileProcessingError) as e:
        print(f"❌ Error de inicialización: {e}")
        return
    except Exception as e:
        logging.error(f"Error inesperado: {e}")
        print(f"❌ Error inesperado: {e}")
        return
    
    # Manejar diferentes modos de ejecución
    if args.stats:
        # Modo estadísticas
        generador.mostrar_dashboard()
        return
    
    if args.test_portales:
        # Modo testing de portales
        print(">>> Generador de CV Inteligente v3.1 - TESTING DE PORTALES")
        generador.test_portales()
        return
    
    if args.debug_html:
        # Modo debug HTML
        print(f">>> Generador de CV Inteligente v3.1 - DEBUG HTML")
        debug_file = generador.debug_html_portal(args.debug_html, "qa")
        if debug_file:
            print(f"\n💡 Para encontrar selectores correctos:")
            print(f"   1. Abre {debug_file} en un navegador")
            print(f"   2. Inspecciona elementos de trabajos")
            print(f"   3. Actualiza config.json con selectores correctos")
        return
    
    if args.scrape:
        # Modo web scraping
        print(">>> Generador de CV Inteligente v3.0 - MODO WEB SCRAPING")
        print(f"🔍 Buscando: {args.scrape}")
        print(f"📍 Ubicación: {args.location}")
        print(f"🕷️ Scraping habilitado: {'SÍ' if generador.config['scraping_config']['enabled'] else 'NO'}\n")
        
        try:
            trabajos = generador.buscar_trabajos_automatico(args.scrape, args.location)
            
            if trabajos:
                # Guardar en CSV si se solicita
                if args.save_jobs:
                    csv_path = generador.guardar_trabajos_csv(trabajos)
                    if csv_path:
                        # Preguntar si procesar automáticamente
                        respuesta = input(f"\n¿Procesar estos {len(trabajos)} trabajos automáticamente? (y/N): ").strip().lower()
                        if respuesta in ['y', 'yes', 'sí', 's']:
                            print(f"\n🚀 Procesando trabajos con modo batch...")
                            resultados = generador.procesar_batch_csv(csv_path)
                            generador.mostrar_resumen_batch(resultados)
                        else:
                            print(f"💾 Trabajos guardados en: {csv_path}")
                            print("💡 Para procesar después: python generador_cv_avanzado.py --batch " + csv_path)
                else:
                    # Solo mostrar resumen
                    print(f"\n📋 TRABAJOS ENCONTRADOS ({len(trabajos)}):")
                    for i, trabajo in enumerate(trabajos[:10], 1):  # Mostrar primeros 10
                        print(f"   {i}. {trabajo['company']} - {trabajo['title']}")
                        if trabajo['salary']:
                            print(f"      💰 {trabajo['salary']}")
                        print(f"      📍 {trabajo['location']} | 🌐 {trabajo['portal']}")
                        print()
                    
                    if len(trabajos) > 10:
                        print(f"   ... y {len(trabajos) - 10} trabajos más")
                    
                    print(f"\n💡 Para guardar y procesar: --scrape {args.scrape} --save-jobs")
            else:
                print("❌ No se encontraron trabajos con los criterios especificados")
                
        except Exception as e:
            print(f"❌ Error en web scraping: {e}")
            logging.error(f"Error en modo scraping: {e}")
        return
    
    if args.batch:
        # Modo batch
        print(">>> Generador de CV Inteligente v3.0 - MODO BATCH")
        print(f"📁 Procesando archivo: {args.batch}")
        print(f"🎯 Umbral mínimo de fit: {generador.umbral_fit}%\n")
        
        try:
            resultados = generador.procesar_batch_csv(args.batch)
            generador.mostrar_resumen_batch(resultados)
        except Exception as e:
            print(f"❌ Error en modo batch: {e}")
        return
    
    if args.empresa and args.postulacion:
        # Modo directo
        print(">>> Generador de CV Inteligente v3.0 - MODO DIRECTO")
        print(f"🏢 Empresa: {args.empresa}")
        print(f"🎯 Umbral mínimo de fit: {generador.umbral_fit}%\n")
        
        try:
            resultado = generador.procesar_postulacion(args.postulacion, args.empresa)
            if resultado:
                print(f"\n✅ ¡Proceso completado para {args.empresa}!")
                print(f">>> Posición: {resultado['titulo']}")
                print(f">>> CV guardado en: {resultado['cv_path']}")
            else:
                print(f"\n❌ Postulación no procesada (fuera de estrategia o fit insuficiente)")
        except Exception as e:
            print(f"❌ Error procesando postulación: {e}")
        return
    
    # Modo interactivo (default)
    print(">>> Generador de CV Inteligente v3.0 - MODO INTERACTIVO")
    print("🎯 Estrategia de aplicación:")
    print("   • Junior: QA, Python, Java, Frontend, Full Stack")
    print("   • Semi-Senior: QA, Python, Java, Full Stack")
    print(f"   • Umbral mínimo de fit: {generador.umbral_fit}%")
    print("   • Rechazo automático: Áreas sin experiencia")
    print("\n💡 Comandos especiales:")
    print("   • 'stats' - Ver dashboard de estadísticas")
    print("   • 'salir' - Terminar programa")
    print("\n🔧 Comandos CLI disponibles:")
    print("   • --batch archivo.csv  - Procesar múltiples postulaciones")
    print("   • --stats              - Ver estadísticas") 
    print("   • --scrape qa --save-jobs - Buscar trabajos automáticamente")
    print("   • --email              - Habilitar envío automático de emails")
    print("   • --help               - Ver ayuda completa")
    
    # Estado del email
    email_status = "✅ HABILITADO" if generador.config['email_config']['enabled'] else "❌ DESHABILITADO"
    print(f"\n📧 Estado de email: {email_status}")
    if generador.config['email_config']['enabled']:
        print(f"   • Configurado: {generador.config['email_config']['email']}")
    else:
        print("   • Para habilitar: actualizar config.json o usar --email")
    print()
    
    while True:
        print("\n" + "="*50)
        empresa = input(">>> Nombre de la empresa (o comando especial): ").strip()
        
        if empresa.lower() == 'salir':
            break
        elif empresa.lower() == 'stats':
            generador.mostrar_dashboard()
            continue
            
        texto_postulacion = input("\n>>> Pega la descripción de la postulación:\n").strip()
        
        if empresa and texto_postulacion:
            try:
                resultado = generador.procesar_postulacion(texto_postulacion, empresa)
                
                if resultado:
                    print(f"\n>>> ¡Proceso completado para {empresa}!")
                    print(f">>> Posición: {resultado['titulo']}")
                    print(f">>> CV guardado en: {resultado['cv_path']}")
                else:
                    print("\n>>> Postulación no procesada (fuera de estrategia o fit insuficiente)")
            except Exception as e:
                logging.error(f"Error procesando postulación de {empresa}: {e}")
                print(f"\n>>> ❌ Error procesando la postulación: {e}")
        else:
            print("\n>>> Empresa y postulación son requeridos")

if __name__ == "__main__":
    main()
